"""
Utilidades para generar y manipular documentos Word y PDF
"""
import os
import subprocess
import logging
import re
import html
import shutil
from pathlib import Path
from typing import Optional, Dict, Any
from docxtpl import DocxTemplate, InlineImage, RichText
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
from io import BytesIO

logger = logging.getLogger(__name__)

# Rutas base para documentos
MEDIA_DIR = Path(__file__).parent.parent.parent / "media"
DOCUMENTOS_DIR = MEDIA_DIR / "documentos"
FIRMAS_DIR = MEDIA_DIR / "firmas"

# Crear directorios si no existen
DOCUMENTOS_DIR.mkdir(parents=True, exist_ok=True)
FIRMAS_DIR.mkdir(parents=True, exist_ok=True)


def _html_a_texto_word(valor: Any) -> Any:
    """
    Convierte HTML simple a texto plano con saltos de línea y viñetas
    para mantener compatibilidad en plantillas docxtpl.
    """
    if not isinstance(valor, str):
        return valor

    contenido = valor.strip()
    if not contenido:
        return ""

    # Si no parece HTML, conservar tal cual (texto plano clásico).
    if not re.search(r"<[^>]+>", contenido):
        return contenido

    # Quitar bloques no imprimibles.
    contenido = re.sub(r"<\s*script[^>]*>.*?<\s*/\s*script\s*>", "", contenido, flags=re.IGNORECASE | re.DOTALL)
    contenido = re.sub(r"<\s*style[^>]*>.*?<\s*/\s*style\s*>", "", contenido, flags=re.IGNORECASE | re.DOTALL)

    # Mapear estructura básica de HTML a texto.
    contenido = re.sub(r"<\s*br\s*/?\s*>", "\n", contenido, flags=re.IGNORECASE)
    contenido = re.sub(r"<\s*/\s*p\s*>", "\n", contenido, flags=re.IGNORECASE)
    contenido = re.sub(r"<\s*p[^>]*>", "", contenido, flags=re.IGNORECASE)
    contenido = re.sub(r"<\s*/\s*div\s*>", "\n", contenido, flags=re.IGNORECASE)
    contenido = re.sub(r"<\s*div[^>]*>", "", contenido, flags=re.IGNORECASE)
    contenido = re.sub(r"<\s*li[^>]*>", "- ", contenido, flags=re.IGNORECASE)
    contenido = re.sub(r"<\s*/\s*li\s*>", "\n", contenido, flags=re.IGNORECASE)
    contenido = re.sub(r"<\s*/\s*(ul|ol)\s*>", "\n", contenido, flags=re.IGNORECASE)
    contenido = re.sub(r"<\s*(ul|ol)[^>]*>", "", contenido, flags=re.IGNORECASE)

    # Eliminar cualquier otra etiqueta remanente.
    contenido = re.sub(r"<[^>]+>", "", contenido)

    # Decodificar entidades HTML y normalizar saltos.
    contenido = html.unescape(contenido)
    contenido = contenido.replace("\r\n", "\n").replace("\r", "\n")
    contenido = re.sub(r"\n{3,}", "\n\n", contenido)

    return contenido.strip()


def _normalizar_contexto_para_word(valores_campos: Dict[str, Any]) -> Dict[str, Any]:
    """
    Devuelve una copia del contexto con valores HTML convertidos a texto imprimible.
    ADEMAS: genera versiones RichText (rt_FIELDNAME) para TODOS los campos.
    Permite usar {{r rt_fieldname}} en plantillas, con o sin HTML en el contenido.
    """
    normalizado: Dict[str, Any] = {}
    for key, value in (valores_campos or {}).items():
        normalizado[key] = _html_a_texto_word(value)

        # Generar RichText automáticamente para TODOS los campos (string o no)
        # Permite usar {{r rt_fieldname}} en plantillas
        normalizado[f"rt_{key}"] = _html_a_richtext_word(value)
    return normalizado


def _html_a_richtext_word(valor: Any) -> Any:
    """
    Convierte HTML básico a RichText de docxtpl para uso con {{r rt_campo}}.
    Fallback: si no hay HTML, devuelve RichText con texto plano.
    """
    if not isinstance(valor, str):
        return valor

    contenido = valor.strip()
    rt = RichText()

    if not contenido:
        return rt

    # Fallback para texto plano
    if not re.search(r"<[^>]+>", contenido):
        rt.add(contenido)
        return rt

    # Limpiar bloques no imprimibles
    contenido = re.sub(r"<\s*script[^>]*>.*?<\s*/\s*script\s*>", "", contenido, flags=re.IGNORECASE | re.DOTALL)
    contenido = re.sub(r"<\s*style[^>]*>.*?<\s*/\s*style\s*>", "", contenido, flags=re.IGNORECASE | re.DOTALL)

    tokens = re.split(r"(<[^>]+>)", contenido)

    bold_level = 0
    italic_level = 0
    underline_level = 0
    list_stack = []
    ol_counters = []

    def _add_text(texto: str) -> None:
        if not texto:
            return
        rt.add(
            texto,
            color="000000",
            bold=bold_level > 0,
            italic=italic_level > 0,
            underline=underline_level > 0
        )

    for token in tokens:
        if not token:
            continue

        if token.startswith("<") and token.endswith(">"):
            tag = token[1:-1].strip().lower()
            if not tag:
                continue

            is_end = tag.startswith("/")
            tag_name = tag[1:].split()[0] if is_end else tag.split()[0]

            if tag_name in ["strong", "b"]:
                if is_end:
                    bold_level = max(0, bold_level - 1)
                else:
                    bold_level += 1
                continue

            if tag_name in ["em", "i"]:
                if is_end:
                    italic_level = max(0, italic_level - 1)
                else:
                    italic_level += 1
                continue

            if tag_name == "u":
                if is_end:
                    underline_level = max(0, underline_level - 1)
                else:
                    underline_level += 1
                continue

            if tag_name == "br":
                _add_text("\n")
                continue

            if tag_name in ["p", "div"] and is_end:
                _add_text("\n")
                continue

            if tag_name == "ul":
                if not is_end:
                    list_stack.append("ul")
                    ol_counters.append(0)
                elif list_stack:
                    list_stack.pop()
                    ol_counters.pop()
                continue

            if tag_name == "ol":
                if not is_end:
                    list_stack.append("ol")
                    ol_counters.append(0)
                elif list_stack:
                    list_stack.pop()
                    ol_counters.pop()
                continue

            if tag_name == "li":
                if not is_end:
                    if list_stack and list_stack[-1] == "ol":
                        ol_counters[-1] += 1
                        _add_text(f"{ol_counters[-1]}. ")
                    else:
                        _add_text("- ")
                else:
                    _add_text("\n")
                continue

            # Etiquetas no soportadas: ignorar
            continue

        texto_plano = html.unescape(token)
        texto_plano = texto_plano.replace("\r\n", "\n").replace("\r", "\n")
        _add_text(texto_plano)

    return rt


def generar_word_desde_plantilla(
    plantilla_path: str,
    documento_id: int,
    valores_campos: Optional[Dict[str, Any]] = None,
    output_filename: Optional[str] = None
) -> str:
    """
    Generar un documento Word desde una plantilla usando docxtpl.
    Reemplaza los placeholders {{campo}} con los valores proporcionados.
    
    Args:
        plantilla_path: Ruta a la plantilla .docx
        documento_id: ID del documento (para nombrar el archivo generado)
        valores_campos: Dict con pares {nombre_campo: valor}
    
    Returns:
        Ruta relativa del documento generado (/static/documentos/...)
    """
    try:
        # Usar valores vacíos si no se proporcionan
        if not valores_campos:
            valores_campos = {}
        
        # Convertir HTML enriquecido a texto compatible con docxtpl.
        valores_campos = _normalizar_contexto_para_word(valores_campos)
        logger.info(f"Generando documento {documento_id} con context: {valores_campos}")
        
        # Cargar plantilla con docxtpl
        doc_template = DocxTemplate(plantilla_path)
        
        # Convertir campos de firma de string → InlineImage
        campos_firma = ['gerente_firma', 'unidad_firma', 'juridica_firma']
        for campo in campos_firma:
            if campo in valores_campos and valores_campos[campo]:
                ruta_firma = valores_campos[campo]
                
                # Convertir /static/firmas/xxx → media/firmas/xxx
                if isinstance(ruta_firma, str) and ruta_firma.startswith('/static/'):
                    ruta_firma = ruta_firma.replace('/static/', '')
                
                # Construir ruta completa
                ruta_completa = MEDIA_DIR / ruta_firma.lstrip('/')
                
                # Si existe el archivo, crear InlineImage
                if ruta_completa.exists():
                    try:
                        valores_campos[campo] = InlineImage(
                            doc_template, 
                            str(ruta_completa), 
                            width=Mm(30)
                        )
                        logger.info(f"Firma {campo} convertida a imagen: {ruta_completa}")
                    except Exception as e:
                        logger.warning(f"Error al cargar firma {campo}: {e}")
                        valores_campos[campo] = ''
                else:
                    logger.warning(f"Archivo de firma no encontrado: {ruta_completa}")
                    valores_campos[campo] = ''
        
        # Renderizar con los valores
        doc_template.render(valores_campos)

        # Reemplazar placeholders en cuadros de texto, parrafos y tablas
        if valores_campos:
            _reemplazar_en_container(doc_template.docx, valores_campos)
            for section in doc_template.docx.sections:
                _reemplazar_en_container(section.header, valores_campos)
                _reemplazar_en_container(section.footer, valores_campos)
        
        # Guardar documento generado
        if not output_filename:
            output_filename = f"{documento_id}_borrador.docx"
        output_path = DOCUMENTOS_DIR / output_filename
        doc_template.save(str(output_path))
        
        # Retornar ruta relativa para servir desde /static
        relative_path = f"/static/documentos/{output_filename}"
        logger.info(f"Documento Word generado: {relative_path}")
        return relative_path
    
    except Exception as e:
        logger.error(f"Error al generar documento Word: {e}")
        raise Exception(f"Error al generar documento: {str(e)}")


def _reemplazar_en_container(container: Any, valores: Dict[str, Any]) -> None:
    """
    Reemplaza placeholders {{campo}} dentro de parrafos, tablas y cuadros de texto.
    """
    try:
        _reemplazar_en_parrafos(container, valores)
        _reemplazar_en_tablas(container, valores)
        _reemplazar_en_textboxes(container, valores)
    except Exception as e:
        logger.warning(f"No se pudo reemplazar texto en el documento: {e}")


def _reemplazar_en_parrafos(container: Any, valores: Dict[str, Any]) -> None:
    for p in getattr(container, "paragraphs", []):
        _reemplazar_en_runs(p, valores)


def _reemplazar_en_tablas(container: Any, valores: Dict[str, Any]) -> None:
    for table in getattr(container, "tables", []):
        for row in table.rows:
            for cell in row.cells:
                _reemplazar_en_parrafos(cell, valores)
                _reemplazar_en_tablas(cell, valores)


def _reemplazar_en_runs(paragraph: Any, valores: Dict[str, Any]) -> None:
    runs = getattr(paragraph, "runs", [])
    if not runs:
        return

    combined = "".join([r.text or "" for r in runs])
    
    # Log para debugging
    if "unidad_nombre" in str(valores) or "gerente" in combined.lower():
        logger.info(f"DEBUG _reemplazar_en_runs: combined='{combined}'")
        logger.info(f"DEBUG _reemplazar_en_runs: claves en valores={list(valores.keys())}")
    
    replacements_made = {}
    for key, value in valores.items():
        placeholder = f"{{{{{key}}}}}"
        # Solo reemplazar texto plano aquí. RichText/InlineImage deben ser
        # procesados por docxtpl en render(), no por reemplazo manual.
        if value is None:
            replacement = ""
        elif isinstance(value, (str, int, float, bool)):
            replacement = str(value)
        else:
            continue
        if placeholder in combined:
            logger.info(f"DEBUG: Reemplazando '{placeholder}' con '{replacement}'")
            combined = combined.replace(placeholder, replacement)
            replacements_made[key] = True

    if replacements_made:
        logger.info(f"DEBUG: Reemplazos realizados en párrafo: {replacements_made}")
        # Solo consolidar runs cuando realmente se reemplazó algo.
        # Si no, mantener runs intactos para no perder formato existente.
        runs[0].text = combined
        for r in runs[1:]:
            r.text = ""


def _reemplazar_en_textboxes(container: Any, valores: Dict[str, Any]) -> None:
    try:
        element = container._element
    except Exception:
        return

    try:
        for txbx in element.xpath(".//w:txbxContent"):
            texts = txbx.xpath(".//w:t")
            if not texts:
                continue

            combined = "".join([t.text or "" for t in texts])

            for key, value in valores.items():
                placeholder = f"{{{{{key}}}}}"
                if value is None:
                    replacement = ""
                elif isinstance(value, (str, int, float, bool)):
                    replacement = str(value)
                else:
                    continue
                if placeholder in combined:
                    combined = combined.replace(placeholder, replacement)

            # Reescribir el contenido consolidado en el primer nodo y vaciar el resto
            texts[0].text = combined
            for t in texts[1:]:
                t.text = ""
    except Exception as e:
        logger.warning(f"No se pudo reemplazar texto en cuadros: {e}")


def incrustar_firma(
    documento_path: str,
    documento_id: int,
    usuario_nombre: str,
    usuario_cargo: str,
    firma_imagen_path: Optional[str] = None
) -> str:
    """
    Incrustar firma (imagen + nombre + cargo) en el documento Word.
    
    Args:
        documento_path: Ruta al documento Word
        documento_id: ID del documento
        usuario_nombre: Nombre del usuario que firma
        usuario_cargo: Cargo del usuario
        firma_imagen_path: Ruta a la imagen de firma (opcional)
    
    Returns:
        Ruta relativa del documento con firma incrustrada
    """
    try:
        # Cargar documento
        doc = Document(documento_path)
        
        # Agregar párrafo de firma
        doc.add_paragraph()  # Espacio
        p_firma = doc.add_paragraph()
        p_firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Agregar imagen de firma si existe
        if firma_imagen_path and os.path.exists(firma_imagen_path):
            try:
                # Redimensionar imagen de firma a 1 inch de ancho
                p_firma.add_run().add_picture(firma_imagen_path, width=Inches(1.5))
                doc.add_paragraph()  # Espacio entre imagen y nombre
            except Exception as e:
                logger.warning(f"No se pudo agregar imagen de firma: {e}")
        
        # Agregar nombre y cargo
        p_nombre = doc.add_paragraph()
        p_nombre.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_nombre = p_nombre.add_run(usuario_nombre)
        run_nombre.bold = True
        run_nombre.font.size = Pt(11)
        
        p_cargo = doc.add_paragraph()
        p_cargo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_cargo = p_cargo.add_run(usuario_cargo)
        run_cargo.font.size = Pt(10)
        run_cargo.font.color.rgb = RGBColor(128, 128, 128)
        
        # Guardar documento con firma
        output_filename = f"{documento_id}_firmado.docx"
        output_path = DOCUMENTOS_DIR / output_filename
        doc.save(str(output_path))
        
        relative_path = f"/static/documentos/{output_filename}"
        logger.info(f"Firma incrustranda en documento: {relative_path}")
        return relative_path
    
    except Exception as e:
        logger.error(f"Error al incrustar firma: {e}")
        raise Exception(f"Error al incrustar firma: {str(e)}")


def convertir_word_a_pdf(
    documento_word_path: str,
    documento_id: int,
    tipo_documento: Optional[str] = None,
    consecutivo: Optional[str] = None
) -> Optional[str]:
    """
    Convertir documento Word a PDF usando LibreOffice.
    
    Args:
        documento_word_path: Ruta al documento Word (.docx)
        documento_id: ID del documento
        tipo_documento: Nombre del tipo de documento (ej: "RESOLUCION", "CIRCULAR")
        consecutivo: Consecutivo asignado (ej: "0220")
    
    Returns:
        Ruta relativa del PDF generado, o None si falla la conversión
    """
    try:
        # Obtener ruta de LibreOffice
        libreoffice_path = _obtener_ruta_libreoffice()
        if not libreoffice_path:
            logger.error("LibreOffice no está instalado o no fue encontrado en el sistema. Rutas buscadas: C:\\Program Files\\LibreOffice, /usr/bin/libreoffice, etc.")
            return None
        
        # Comando para convertir a PDF
        output_dir = str(DOCUMENTOS_DIR)
        cmd = [
            libreoffice_path,
            "--headless",
            "--convert-to", "pdf",
            "--outdir", output_dir,
            documento_word_path
        ]
        
        logger.info(f"Ejecutando conversión a PDF con comando: {' '.join(cmd)}")
        
        # Ejecutar conversión
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        
        if result.returncode != 0:
            logger.error(f"Error en conversión a PDF. Código: {result.returncode}")
            logger.error(f"STDOUT: {result.stdout}")
            logger.error(f"STDERR: {result.stderr}")
            return None
        
        # Obtener nombre del PDF generado
        word_filename = Path(documento_word_path).stem
        pdf_filename = f"{word_filename}.pdf"
        pdf_path = DOCUMENTOS_DIR / pdf_filename
        
        if not pdf_path.exists():
            logger.warning(f"PDF no generado: {pdf_path}")
            return None
        
        # Renombrar a formato estándar: TIPO_DOCUMENTO_N°_CONSECUTIVO.pdf
        if tipo_documento and consecutivo:
            # Limpiar nombre del tipo para archivo
            tipo_limpio = tipo_documento.upper().replace(" ", "_")
            final_pdf_name = f"{tipo_limpio}_N°_{consecutivo}.pdf"
        else:
            final_pdf_name = f"{documento_id}_final.pdf"
        
        final_pdf_path = DOCUMENTOS_DIR / final_pdf_name
        # Usar shutil.move() que maneja archivos existentes correctamente (en Windows Path.rename() falla)
        shutil.move(str(pdf_path), str(final_pdf_path))
        
        relative_path = f"/static/documentos/{final_pdf_name}"
        logger.info(f"PDF generado: {relative_path}")
        return relative_path
    
    except subprocess.TimeoutExpired:
        logger.error("Timeout en conversión a PDF")
        return None
    except Exception as e:
        logger.error(f"Error al convertir Word a PDF: {e}")
        return None


def _obtener_ruta_libreoffice() -> Optional[str]:
    """
    Obtener ruta del ejecutable de LibreOffice según el SO.
    """
    import sys
    import os

    # Permitir sobreescribir la ruta por variable de entorno
    env_path = os.getenv("LIBREOFFICE_PATH")
    if env_path and os.path.exists(env_path):
        logger.info(f"LibreOffice encontrado en variable LIBREOFFICE_PATH: {env_path}")
        return env_path
    
    if sys.platform == "win32":
        # Rutas comunes en Windows (agregadas más variantes)
        rutas_posibles = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            r"C:\LibreOffice\program\soffice.exe",
            r"C:\Program Files\LibreOffice 7\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice 7\program\soffice.exe",
        ]
    elif sys.platform == "darwin":
        # macOS
        rutas_posibles = [
            "/Applications/LibreOffice.app/Contents/MacOS/soffice"
        ]
    else:
        # Linux
        rutas_posibles = [
            "/usr/bin/libreoffice",
            "/usr/bin/soffice",
            "/usr/local/bin/libreoffice"
        ]
    
    for ruta in rutas_posibles:
        if os.path.exists(ruta):
            logger.info(f"LibreOffice encontrado en: {ruta}")
            return ruta
    
    logger.warning(f"LibreOffice no encontrado en rutas esperadas: {rutas_posibles}")
    
    # Intentar encontrar en PATH
    try:
        cmd = ["which", "soffice"] if sys.platform != "win32" else ["where", "soffice.exe"]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=5)
        if result.returncode == 0:
            ruta_encontrada = result.stdout.strip()
            logger.info(f"LibreOffice encontrado en PATH: {ruta_encontrada}")
            return ruta_encontrada
    except Exception as e:
        logger.warning(f"No se pudo buscar LibreOffice en PATH: {e}")
    
    logger.error("LibreOffice no está instalado o no fue encontrado en el sistema")
    return None


def guardar_firma_usuario(usuario_id: int, imagen_bytes: bytes) -> str:
    """
    Guardar imagen de firma de usuario.
    
    Args:
        usuario_id: ID del usuario
        imagen_bytes: Bytes de la imagen
    
    Returns:
        Ruta relativa de la firma guardada
    """
    try:
        # Guardar imagen
        filename = f"firma_usuario_{usuario_id}.png"
        filepath = FIRMAS_DIR / filename
        
        with open(filepath, 'wb') as f:
            f.write(imagen_bytes)
        
        relative_path = f"/static/firmas/{filename}"
        logger.info(f"Firma guardada: {relative_path}")
        return relative_path
    
    except Exception as e:
        logger.error(f"Error al guardar firma: {e}")
        raise Exception(f"Error al guardar firma: {str(e)}")


def eliminar_archivo_documento(documento_id: int, ruta_relativa: Optional[str] = None) -> bool:
    """
    Eliminar archivo .docx de un documento de la carpeta media/documentos.
    Si no se proporciona ruta_relativa, intenta eliminar los archivos conocidos.
    
    Args:
        documento_id: ID del documento
        ruta_relativa: Ruta relativa del archivo (/static/documentos/...)
        
    Returns:
        True si se eliminó, False si no encontró archivo
    """
    try:
        rutas_a_intentar = []
        
        # Si se proporciona ruta, intentar eliminarla
        if ruta_relativa:
            # Convertir ruta web a ruta física
            if ruta_relativa.startswith('/static/documentos/'):
                filename = ruta_relativa.replace('/static/documentos/', '')
                rutas_a_intentar.append(DOCUMENTOS_DIR / filename)
            else:
                rutas_a_intentar.append(Path(ruta_relativa))
        
        # También intentar patrones comunes
        rutas_a_intentar.extend([
            DOCUMENTOS_DIR / f"{documento_id}_borrador.docx",
            DOCUMENTOS_DIR / f"{documento_id}_borrador_firmado.docx",
            DOCUMENTOS_DIR / f"{documento_id}_final.docx",
            DOCUMENTOS_DIR / f"doc_{documento_id}.docx",
        ])
        
        eliminados = False
        for filepath in rutas_a_intentar:
            try:
                if filepath.exists() and filepath.is_file():
                    filepath.unlink()
                    logger.info(f"Archivo eliminado: {filepath}")
                    eliminados = True
            except Exception as e:
                logger.warning(f"No se pudo eliminar {filepath}: {e}")
        
        return eliminados
    
    except Exception as e:
        logger.error(f"Error al eliminar archivo de documento {documento_id}: {e}")
        return False
